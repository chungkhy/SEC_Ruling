import docx
import pandas as pd
import nltk
import re
import json
import openai
from openai import OpenAI


class SEC_Report_Reference_Extractor:
    def __init__(self, docx_file_path, sec_data_excel_path, api_key):
        """
        Initialize the extractor with the document and data paths.
        """
        self.docx_file_path = docx_file_path
        self.sec_data_excel_path = sec_data_excel_path
        self.api_key = api_key
        self.document = docx.Document(self.docx_file_path)

        # Get the complete text from the document
        self.combined_text = ""
        self.get_complete_text()

        # Initialize OpenAI client
        openai.api_key = self.api_key
        self.client = OpenAI(api_key=self.api_key)

        # Load SEC data and extract the list of commenter names
        self.sec_data = pd.read_excel(
            self.sec_data_excel_path, sheet_name="Organizations_People_Initiative"
        )
        self.commenters_only = self.sec_data[
            self.sec_data["Role viewed from SEC Perspective_1"] == "commenter"
        ].copy()
        self.commenters_only["Other Names"].fillna(
            self.commenters_only["Name"], inplace=True
        )
        self.name_list = self.commenters_only["Other Names"].unique()

        # Download NLTK data files required for sentence tokenization
        nltk.download("punkt", quiet=True)
        nltk.download('punkt_tab')

        # Initialize variables to store data
        self.sentences_with_superscript = []  # List of sentences with superscript references: (sentence, [reference_numbers])
        self.reference_numbers = set()        # Set of all reference numbers found
        self.collected_references = set()     # Set of reference numbers successfully extracted
        self.missing_references = set()       # Set of reference numbers not extracted successfully
        self.store_responses = {}

        # DataFrame to store the final results
        self.successful_data = pd.DataFrame()

    def get_complete_text(self):
        """
        Get the complete text from the original document.
        """
        for para in self.document.paragraphs:

            # Skip the empty paragraphs and footnotes
            if para.text.strip() == "" or any(run.font.superscript and run.text.isdigit() for run in para.runs):
                continue
            self.combined_text += para.text.strip() + "\n"

    def extract_superscript_references(self):
        """
        Extract superscript reference numbers and their corresponding sentences from the document.
        """
        print("Extracting superscript references...")
        full_text = ""

        # Iterate over paragraphs and runs to build the full text with markers for references
        for para in self.document.paragraphs:
            for run in para.runs:
                text = run.text
                if run.font.superscript and text.isdigit():

                    # Replace superscript number with a unique marker
                    ref_number = int(text)
                    marker = f"[{ref_number}]"
                    full_text += marker
                    self.reference_numbers.add(ref_number)
                else:
                    full_text += text

            # Add a newline after each paragraph for clarity
            full_text += "\n"

        # Split the full text into sentences
        sentences = nltk.sent_tokenize(full_text)

        # Compile a regex pattern to find reference numbers in brackets
        ref_pattern = re.compile(r'\[(\d+)\]')

        # Iterate over each sentence to find and extract references
        for sentence in sentences:
            matches = ref_pattern.findall(sentence)
            if matches:
                ref_numbers = [int(num) for num in matches]

                # Remove the reference numbers from the sentence for clean text
                clean_sentence = ref_pattern.sub('', sentence).strip()
                self.sentences_with_superscript.append((clean_sentence, ref_numbers))
        print("Superscript references extracted.")

    def extract_with_openai(self):
        """
        Use OpenAI API to extract references and its corresponding sentences.
        """
        print("Extracting references with OpenAI...")

        # Prepare the user prompt for GPT
        user_prompt = (
            "In the following text, extract each reference number and the corresponding sentence. "
            "Each reference number will only map to one sentence. "
            "For each sentence, classify it into one of the following categories: 'Agree', 'Disagree', or 'Neutral' "
            "based on its content. "
            "Output the result as a JSON array where each element has 'Reference': reference_number, "
            "'Sentence': sentence, and 'Classification': classification.\n\n"
            f"Text:\n{self.combined_text}"
        )

        # Make API call to OpenAI's GPT model
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an assistant who extracts references, their corresponding sentences, and classifies them from the text.",
                    },
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.7,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0,
                response_format={"type": "json_object"}
            )

            # Store the response content as a dictionary
            self.store_responses = json.loads(response.choices[0].message.content)
            print("Received response:\n", self.store_responses)
            print("Extraction complete.")

        # Handle exceptions
        except json.JSONDecodeError as json_err:
            print(f"Error parsing JSON response: {json_err}")
        except Exception as e:
            print(f"An error occurred during OpenAI API call: {e}")

    def process_responses(self):
        """
        Process the responses from OpenAI and handle any issues to ensure all references are extracted.
        """
        print("Processing GPT responses...")

        # Remove the extra layers if present
        if "Sentences" in self.store_responses:
            self.store_responses = self.store_responses["Sentences"]
        elif "results" in self.store_responses:
            self.store_responses = self.store_responses["results"]
        elif "result" in self.store_responses:
            self.store_responses = self.store_responses["result"]
        elif "References" in self.store_responses:
            self.store_responses = self.store_responses["References"]

        # Extract sentences and references
        extracted_data = []
        for sentence_data in self.store_responses:
            # Check if the response contains the expected keys
            if not ("Sentence" in sentence_data and "Reference" in sentence_data and "Classification" in sentence_data):
                continue

            # Extract the sentence, reference, and classification data
            reference = int(sentence_data["Reference"])
            sentence = sentence_data["Sentence"]
            classification = sentence_data["Classification"]
            self.collected_references.add(reference)
            extracted_data.append([reference, sentence, classification])

        # Create a DataFrame from the extracted data
        self.successful_data = pd.DataFrame(extracted_data, columns=["Reference", "Sentence", "GPT_Classification"])

        # Re-extract any missing references
        self.missing_references = self.reference_numbers - self.collected_references
        if self.missing_references:
            self.redo_extracting_and_processing()
        else:
            print("Processing complete without missing references.")

    def redo_extracting_and_processing(self):
        """
        Re-extract the missing references and re-process the sentences with issues.
        """
        print("Re-extracting and re-processing missing references...")
        print(f"Missing references: {sorted(self.missing_references)}")

        # Re-extract the missing references using OpenAI
        store_responses_redo = {}
        user_prompt = (
            f"In the following text, extract the following reference numbers: {sorted(self.missing_references)}, "
            "and their corresponding sentences. Each reference number will only map to one sentence. "
            "For each sentence, classify it into one of the following categories: 'Agree', 'Disagree', or 'Neutral' "
            "based on its content. "
            "Output the result as a JSON array where each element has 'Sentence': sentence, "
            "'Reference': reference_number, and 'Classification': classification.\n\n"
            f"Text:\n{self.combined_text}"
        )

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an assistant who extracts references, their corresponding sentences, and classifies them from the text.",
                    },
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.7,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0,
                response_format={"type": "json_object"}
            )

            store_responses_redo = json.loads(response.choices[0].message.content)

            # Remove the extra layers if present
            if "Sentences" in self.store_responses:
                self.store_responses = self.store_responses["Sentences"]
            elif "results" in self.store_responses:
                self.store_responses = self.store_responses["results"]
            elif "result" in self.store_responses:
                self.store_responses = self.store_responses["result"]
            elif "References" in self.store_responses:
                self.store_responses = self.store_responses["References"]

            # Extract sentences, references, and classifications
            extracted_data_redo = []
            for sentence_data in store_responses_redo:
                if not ("Sentence" in sentence_data and "Reference" in sentence_data and "Classification" in sentence_data):
                    continue
                reference = int(sentence_data["Reference"])
                sentence = sentence_data["Sentence"]
                classification = sentence_data["Classification"]
                self.collected_references.add(reference)
                extracted_data_redo.append([reference, sentence, classification])

            # Concatenate the new data with the existing DataFrame
            new_data = pd.DataFrame(extracted_data_redo, columns=["Reference", "Sentence", "GPT_Classification"])
            self.successful_data = pd.concat([self.successful_data, new_data], ignore_index=True)
            print("Re-extraction and re-processing complete.")

        # Handle exceptions
        except json.JSONDecodeError as json_err:
            print(f"Error parsing JSON response: {json_err}")
        except Exception as e:
            print(f"An error occurred during OpenAI API call: {e}")

    # Classification (OpenAI and BERT or RoBERTa model): Agree, Disagree, Neutral

    def save_results(self, output_excel_path):
        """
        Save the final results to an Excel file.
        """
        # Sort the DataFrame by 'Reference' column for better organization
        self.successful_data.sort_values(by="Reference", inplace=True)
        self.successful_data.to_excel(output_excel_path, index=False)
        print(f"Results saved to {output_excel_path}")

    def run(self, output_excel_path):
        """
        Run the extraction process and save the results.
        """
        self.extract_superscript_references()
        self.extract_with_openai()
        self.process_responses()
        self.save_results(output_excel_path)


# Example usage
if __name__ == "__main__":
    docx_file_path = "33-11275_C1_short.docx"
    sec_data_excel_path = "SEC Data.xlsx"
    api_key = "<API KEY>"
    output_excel_path = "extracted_references.xlsx"

    # Create and run the extractor
    extractor = SEC_Report_Reference_Extractor(docx_file_path, sec_data_excel_path, api_key)
    extractor.run(output_excel_path)
